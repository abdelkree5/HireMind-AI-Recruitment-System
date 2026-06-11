"""Create 5 test CVs as .docx files for testing."""
import os
from docx import Document

os.makedirs("test_cvs", exist_ok=True)

cvs = {
    "ahmed_backend.docx": {
        "name": "Ahmed Mohamed",
        "sections": [
            ("", "Ahmed Mohamed\nFull Stack / Backend Developer\nLocation: Cairo, Egypt\nEmail: ahmed@example.com"),
            ("Summary", "Passionate backend developer with 4 years of experience building scalable microservices and APIs using Python, FastAPI, and PostgreSQL."),
            ("Experience", "Senior Backend Engineer | Tech Solutions\n- Designed and developed high-performance REST APIs using FastAPI and Django.\n- Managed PostgreSQL database optimizations and migrations.\n- Implemented background tasks using Celery and Redis.\n- Dockerized applications and set up CI/CD pipelines."),
            ("Skills", "Python, FastAPI, Django, PostgreSQL, Docker, Redis, REST APIs, Git, Linux"),
        ]
    },
    "sarah_frontend.docx": {
        "name": "Sarah Ahmed",
        "sections": [
            ("", "Sarah Ahmed\nFrontend Web Developer\nEmail: sarah@example.com"),
            ("Summary", "Frontend developer with 2 years of experience in designing and developing modern, responsive web applications using React and Vite."),
            ("Experience", "Frontend Developer | Web Soft\n- Developed user interfaces using React.js and TailwindCSS.\n- Consumed APIs and integrated them with the interactive UI.\n- Optimized application performance for speed across all devices.\n- Wrote tests using Jest and RTL."),
            ("Skills", "React, JavaScript, HTML, CSS, TailwindCSS, Vite, Redux, Git"),
        ]
    },
    "omar_data_scientist.docx": {
        "name": "Omar Ali",
        "sections": [
            ("", "Omar Ali\nData Scientist & Machine Learning Engineer\nEmail: omar@example.com"),
            ("Profile", "Data Scientist with a master's degree in Computer Science. Proficient in developing predictive models and analyzing large datasets."),
            ("Experience", "Data Scientist | Data Minds Inc\n- Built machine learning models for customer churn prediction using Scikit-Learn and XGBoost.\n- Performed NLP tasks including sentiment analysis with HuggingFace Transformers and PyTorch.\n- Analyzed data using Pandas, NumPy, and SQL.\n- Visualized results with Matplotlib and Tableau."),
            ("Skills", "Python, Machine Learning, PyTorch, Scikit-Learn, NLP, Transformers, Pandas, SQL, Data Visualization"),
        ]
    },
    "mahmoud_devops.docx": {
        "name": "Mahmoud Hassan",
        "sections": [
            ("", "Mahmoud Hassan\nDevOps & Cloud Engineer\nEmail: mahmoud@example.com"),
            ("Summary", "Experienced DevOps engineer focused on cloud infrastructure, CI/CD, and automation."),
            ("Experience", "Cloud Engineer | CloudTech\n- Provisioned infrastructure on AWS using Terraform.\n- Managed Kubernetes clusters and deployed microservices.\n- Automated testing and deployment workflows via GitHub Actions.\n- Monitored systems with Prometheus and Grafana."),
            ("Skills", "AWS, Kubernetes, Docker, Terraform, CI/CD, GitHub Actions, Linux, Bash, Prometheus"),
        ]
    },
    "nour_mobile.docx": {
        "name": "Nour Abdelkreem",
        "sections": [
            ("", "Nour Abdelkreem\nMobile App Developer (Flutter)\nEmail: nour@example.com"),
            ("Objective", "Mobile application developer specialized in Flutter, striving to develop engaging applications for both Android and iOS platforms."),
            ("Experience", "Mobile Developer | Appify\n- Developed cross-platform applications using the Flutter framework and Dart.\n- Integrated applications with Firebase databases.\n- Managed application state using Provider and Riverpod.\n- Deployed applications to Google Play and the App Store."),
            ("Skills", "Flutter, Dart, Firebase, RESTful APIs, UI/UX, Git"),
        ]
    },
}

for filename, cv_data in cvs.items():
    doc = Document()
    for heading, content in cv_data["sections"]:
        if heading:
            doc.add_heading(heading, level=1)
        for line in content.split("\n"):
            doc.add_paragraph(line)
    doc.save(os.path.join("test_cvs", filename))
    print(f"Created: {filename}")

print("\nAll 5 DOCX CVs created successfully!")
